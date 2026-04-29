import os
import sys
import uuid
import shutil
from flask import Flask, render_template, request, jsonify, send_file
import fitz
from PIL import Image as PILImage, ImageOps
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import concurrent.futures
import zipfile

# Detectar si estamos en modo congelado (PyInstaller)
if getattr(sys, 'frozen', False):
    # Estamos en el ejecutable
    base_path = sys._MEIPASS
else:
    base_path = os.path.dirname(os.path.abspath(__file__))

app = Flask(__name__, static_folder=os.path.join(os.path.dirname(base_path), 'static'), static_url_path='/static')
app.config['UPLOAD_FOLDER'] = os.path.join(base_path, 'uploads')
app.config['OUTPUT_FOLDER'] = os.path.join(base_path, 'outputs')
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

def convert_page_worker(args):
    pdf_path, page_num, output_file, dpi, quality, auto_crop = args
    try:
        doc = fitz.open(pdf_path)
        page = doc.load_page(page_num)
        zoom = dpi / 72
        rect = page.rect
        width, height = rect.width * zoom, rect.height * zoom
        MAX_DIM = 8000
        if width > MAX_DIM or height > MAX_DIM:
            scale = MAX_DIM / max(width, height)
            zoom *= scale
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img = PILImage.frombytes("RGB", (pix.width, pix.height), pix.samples)
        if auto_crop:
            gray_im = img.convert("L")
            invert_im = ImageOps.invert(gray_im)
            mask = invert_im.point(lambda x: 0 if x < 5 else 255)
            bbox = mask.getbbox()
            if bbox:
                left, top, right, bottom = bbox
                left = max(0, left - 15)
                top = max(0, top - 15)
                right = min(img.width, right + 15)
                bottom = min(img.height, bottom + 15)
                img = img.crop((left, top, right, bottom))
        if img.width < 10 or img.height < 10:
            doc.close()
            return False
        img.save(output_file, "JPEG", quality=quality, optimize=True)
        doc.close()
        return True
    except Exception as e:
        print(f"Error en convert_page_worker: {e}")
        return False

def generate_word(images, output_path):
    try:
        doc = Document()
        images.sort()
        for i, img_path in enumerate(images):
            if i > 0:
                section = doc.add_section()
            else:
                section = doc.sections[0]
            with PILImage.open(img_path) as img:
                w_px, h_px = img.size
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            img_w_in = w_px / 96
            img_h_in = h_px / 96
            scale = min(1.0, 7.5 / img_w_in, 10.0 / img_h_in)
            final_w = img_w_in * scale
            final_h = img_h_in * scale
            section.top_margin = Inches((11.0 - final_h) / 2)
            section.left_margin = section.right_margin = Inches(0.5)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(img_path, width=Inches(final_w))
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error en generate_word: {e}")
        return False

@app.route('/')
def index():
    return send_file('../static/index.html')

@app.route('/upload', methods=['POST'])
def upload():
    files = request.files.getlist('files')
    mode = request.form.get('mode', 'pdf')
    task_id = str(uuid.uuid4())
    task_folder = os.path.join(app.config['UPLOAD_FOLDER'], task_id)
    os.makedirs(task_folder, exist_ok=True)
    for f in files:
        f.save(os.path.join(task_folder, f.filename))
    return jsonify({'task_id': task_id, 'mode': mode})

@app.route('/process', methods=['POST'])
def process():
    data = request.json
    task_id = data['task_id']
    mode = data['mode']
    dpi = {'Estandar': 150, 'HD': 300, 'Ultra': 600}.get(data.get('quality', 'HD'), 300)
    
    task_folder = os.path.join(app.config['UPLOAD_FOLDER'], task_id)
    output_folder = os.path.join(app.config['OUTPUT_FOLDER'], task_id)
    os.makedirs(output_folder, exist_ok=True)
    
    word_files = []
    
    if mode == 'pdf':
        pdfs = [f for f in os.listdir(task_folder) if f.lower().endswith('.pdf')]
        tasks = []
        for pdf in pdfs:
            pdf_path = os.path.join(task_folder, pdf)
            name = os.path.splitext(pdf)[0]
            doc = fitz.open(pdf_path)
            for i in range(len(doc)):
                out = os.path.join(output_folder, f"{name}_p{i+1}.jpg")
                tasks.append((pdf_path, i, out, dpi, 95, True))
            doc.close()
        
        # Procesar todas las páginas a la vez (rápido)
        with concurrent.futures.ProcessPoolExecutor(max_workers=os.cpu_count()) as executor:
            results = list(executor.map(convert_page_worker, tasks))
        
        if data.get('gen_word', True):
            imgs = [os.path.join(output_folder, f) for f in os.listdir(output_folder) if f.endswith('.jpg')]
            word_path = os.path.join(output_folder, "Imagenes.docx")
            if generate_word(imgs, word_path):
                word_files.append(word_path)
    else:
        imgs = [os.path.join(task_folder, f) for f in os.listdir(task_folder) if f.lower().endswith(('.jpg', '.jpeg', '.png'))]
        if imgs:
            word_path = os.path.join(output_folder, "Imagenes.docx")
            if generate_word(imgs, word_path):
                word_files.append(word_path)
    
    if len(word_files) > 0:
        return jsonify({'ready': True, 'task_id': task_id})
    else:
        return jsonify({'ready': False, 'error': 'No se generaron archivos'})

@app.route('/download/<task_id>/<filename>')
def download(task_id, filename):
    path = os.path.join(app.config['OUTPUT_FOLDER'], task_id, filename)
    upload_path = os.path.join(app.config['UPLOAD_FOLDER'], task_id)
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], task_id)
    
    response = send_file(path, as_attachment=True)
    
    @response.call_on_close
    def cleanup():
        try:
            if os.path.exists(upload_path):
                shutil.rmtree(upload_path)
            if os.path.exists(output_path):
                shutil.rmtree(output_path)
        except Exception as e:
            print(f"Error en limpieza: {e}")
    
    return response

if __name__ == '__main__':
    # Detectar si estamos en modo congelado (PyInstaller)
    if getattr(sys, 'frozen', False):
        # Estamos en el ejecutable
        os.chdir(base_path)
    
    port = 5000
    print(f"Running on http://localhost:{port}/")
    app.run(debug=False, port=port)
